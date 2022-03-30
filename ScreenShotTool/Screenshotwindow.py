from tkinter import *
import os
import sys
import time
import tkinter.messagebox as tmsg
from PIL import Image, ImageTk
import pyautogui
from tkinter import ttk
from threading import Thread
from get_window import get_title
from forground import trigger
from createDocument import create_document
from floating_button import floating_window
from datetime import datetime
from byactivewindow import takeScreenshot
from docx import Document
from docx.shared import Inches, Mm
home=None
QuickSNAP=None
refresh_button=None
startip_flag = 1
active_title=""
image_name = "startup.png"
Message=""
path=""


def message_show(Message):
    message_label.config(text=Message)


def mainfunction(root, working_dir, document_name):
    global path
    path=working_dir+"\\{}".format(document_name)
    window1 = Toplevel()
    window1.iconbitmap("main_icon.ico")
    window1.geometry("1370x720")
    window1.minsize(1370, 720)
    # window1.maxsize(1330, 710)
    window1.title("ScreenShotTK {}".format(path))
    title="default"

    ################ CREATING FRAME ************************

    f1 = Frame(window1,borderwidth=6, relief=SUNKEN,width=100,background="#7092bc")
    f1.pack(side=LEFT, fill='both')

    f2 = Frame(window1, relief=SUNKEN)
    f2.pack(side=TOP, fill="x")

    f3 = Frame(window1, relief=SUNKEN)
    f3.pack(fill="x")

    f4 = Frame(window1,  relief=SUNKEN)
    f4.pack(side=BOTTOM, fill="x")

    def read_metadata():
        global metadata
        f = open("metadata.txt", "r")
        metadata = f.read()
        if metadata.startswith("Enter metadata which you"):
            metadata=""

    def image_handle():
        global photo
        path_image=os.getcwd()
        image = Image.open(path_image+"\\startup.png")
        image = image.resize((800, 500), Image.ANTIALIAS)
        photo = ImageTk.PhotoImage(image)

        global close_photo
        close_image =  Image.open(path_image+"\\home-page.png")
        close_image = close_image.resize((50, 50), Image.ANTIALIAS)
        close_photo = ImageTk.PhotoImage(close_image)

        global imagefolder_icon
        image_folder = Image.open(path_image+"\\pictures-folder.png")
        image_folder = image_folder.resize((50, 50), Image.ANTIALIAS)
        imagefolder_icon = ImageTk.PhotoImage(image_folder)

        global docsfolder_icon
        doc_folder = Image.open(path_image+"\\documents-folder.png")
        doc_folder = doc_folder.resize((50, 50), Image.ANTIALIAS)
        docsfolder_icon = ImageTk.PhotoImage(doc_folder)

        global screenshot_icon
        screenshot = Image.open(path_image+"\\screenshot.png")
        screenshot = screenshot.resize((50, 50), Image.ANTIALIAS)
        screenshot_icon = ImageTk.PhotoImage(screenshot)

        global start_icon
        starticon = Image.open(path_image + "\\start.png")
        starticon = starticon.resize((50, 50), Image.ANTIALIAS)
        start_icon = ImageTk.PhotoImage(starticon)

        global stop_icon
        stopicon = Image.open(path_image + "\\pause.png")
        stopicon = stopicon.resize((50, 50), Image.ANTIALIAS)
        stop_icon = ImageTk.PhotoImage(stopicon)

        global save_icon
        saveicon = Image.open(path_image + "\\save.png")
        saveicon = saveicon.resize((50, 50), Image.ANTIALIAS)
        save_icon = ImageTk.PhotoImage(saveicon)


    # ********* Header Section **********************************
    def headerSection():

        #  ******************* Refresh Button ************************
        def refreshbutton():
            global refresh_button, active_title
            refresh_button=ttk.Button(f2, text="Refresh", command=headerSection)
            refresh_button.grid(column=3, row=0, padx=10, pady=25)
            active_title=None

        refreshbutton()


        #      First Grid                             #
        Label(f2, text="Select Active Windows",
                  font=("Times New Roman", 12, "bold")).grid(column=1,
                                                     row=0, padx=10, pady=25)
        global title
        title=list(set(get_title()))

        # ***** COMBO BOX **********
        def comboclick(event):
            global active_title
            active_title = active_window.get()

        global active_window
        active_window = ttk.Combobox(f2, width=50,
                                         value=title)
        active_window.grid(column=2, row=0)
        active_window.current(0)
        active_window.bind('<<ComboboxSelected>>', comboclick)


    # ***************** SIDE SECTION *****************************

    def side_section():
        color={"frame_bg":"#7092bc", "time_label_bg":"black","highlightbackground":"white", "highlightcolor":"red"}
        f0 = Frame(f1, relief=SUNKEN, bg=color["frame_bg"])
        f0.pack(fill="both")

            # *************** INTERVAL SETUP *****************************
        Label(f0, text="Setup",  font="serif 10", bg=color["frame_bg"]).grid(row=3,sticky=W, pady=10)
        f5 = Frame(f1, relief=SUNKEN, bg=color["frame_bg"] ,highlightbackground=color["highlightbackground"], highlightcolor=color["highlightcolor"], highlightthickness=1)
        f5.pack(fill="both", pady=5)
        Label(f5, text="Set Interval", font="serif 10",bg=color["frame_bg"] ).grid(row=4, sticky=W, pady=10)

        hour=IntVar()
        Spinbox(f5,from_=0, to=23,text= hour,font="serif 10 bold",bg=color["frame_bg"], width=15).grid(row=6, column=0, sticky=W, padx=5)
        Label(f5, text="HH", height=1 ,font="serif 10 bold",bg=color["frame_bg"],fg="black").grid(row=6, column=1,padx=5)

        minute = IntVar()
        Spinbox(f5, from_=0, to=59,text= minute, font="serif 10 bold",bg=color["frame_bg"], width=15).grid(row=7, column=0, sticky=W, padx=5)
        Label(f5, text="MM", height=1 ,font="serif 10 bold",bg=color["frame_bg"],fg="black").grid(row=7, column=1, padx=5)

        second = StringVar()
        Spinbox(f5, from_=0, to=59,text= second, font="serif 10 bold",bg=color["frame_bg"], width=15).grid(row=8, column=0, sticky=W, padx=5)
        Label(f5, text="SEC",height=1 ,font="serif 10 bold",bg=color["frame_bg"],fg="black").grid(row=8, column=1, padx=5)

        milisec = StringVar()
        Spinbox(f5, from_=0, to=59,text= milisec, font="serif 10 bold",bg=color["frame_bg"], width=15).grid(row=9, column=0, sticky=W, padx=5)
        Label(f5, text="MS", height=1, font="serif 10 bold",bg=color["frame_bg"],fg="black").grid(row=9, column=1, padx=5)
        Label(f5, text="", height=1, font="serif 8",bg=color["frame_bg"],fg="white", wraplength=15).grid(row=10, column=1, padx=5)

            # ******************** SELECT ACTION TYPE-- ByActive Window OR ByTitle *****************
        f6 = Frame(f1, relief=SUNKEN, bg=color["frame_bg"], pady=12 ,highlightbackground=color["highlightbackground"], highlightcolor=color["highlightcolor"], highlightthickness=1)
        f6.pack(fill="x")

        def deactive(value):

            if value == 1:
                active_window.config(state='disabled')
                refresh_button.config(state='disabled')

            if value == 2:
                active_window.config(state='normal')
                refresh_button.config(state='normal')


        def call_screenshot():
            global execution, Message
            sleep_time = calculate_time()
            if sleep_time != -1:
                stop.config(state='normal', bg=bg_button["color"])
                start.config(state='disabled', bg="white")
                save_button.config(state='disabled', bg=bg_button["color"])
                while not execution :
                    if action_type.get()== 1:
                        Message=takeScreenshot(path)
                        message_show(Message)
                    else:
                        Message=trigger(active_title, path)
                        message_show(Message)
                    time.sleep(sleep_time)
                return

        def thread_create():
            global execution
            execution=False
            global home, QuickSNAP
            home.config(state="disabled")
            QuickSNAP.config(state="disabled")
            t = Thread(target=call_screenshot)
            t.start()


        def calculate_time():

            hr=hour.get()
            min=minute.get()
            sec=second.get()
            ms=milisec.get()

            if hr !="" and min !="" and  sec !="" and ms !="":
                if (int(hr) >=0) and  (int(min) >=0 ) and (int(sec) >=0 ) and (int(ms) >=0 ):
                    return int(hr) *3600 + int(min) * 60 + int(sec) + (int(ms)/1000)%60
                else:
                    tmsg.showerror("Warning", "Enter valid interval", parent=window1)
                    return -1
            else:
                tmsg.showerror("Warning","Enter valid interval", parent=window1)
                return  -1

        action_type=IntVar()
        label=Label(f6, text="Select Action Type", font="Helvetica 10", justify=LEFT, bg=color["frame_bg"])
        label.grid(sticky=W, pady=5)


 #***************   RADIO BUTTON *****************************
        r1=Radiobutton(f6, text="ByActive Window(Recommend)", variable= action_type,bg=color["frame_bg"],  value=1 ,font="comicsansms 10", command=lambda: deactive(action_type.get()))
        r1.grid(row=2, column=0, pady=10, padx=10, sticky=W)

        r2=Radiobutton(f6, text="ByTitle", variable= action_type, value=2 ,bg=color["frame_bg"],font="comicsansms 10", command=lambda: deactive(action_type.get()))
        r2.grid(row=3, column=0, pady=10,padx=10, sticky=W)
        r2.select()

        # **************** START, STOP AND SAVE BUTTON ***********************
        def quit_():
            start.config(state='normal',bg=bg_button["color"])
            stop.config(state='disabled', bg=bg_button["color"])
            save_button.config(state='normal',bg=bg_button["color"])
            global execution
            execution=True
            global home, QuickSNAP
            home.config(state="normal")
            QuickSNAP.config(state="normal")

        bg_button={"color":"#6699ff"}
        def save():
            start.config(state='disabled', bg=bg_button["color"])
            stop.config(state='disabled', bg=bg_button["color"])
            save_button.config(state='disabled', bg="white")
            global home, QuickSNAP
            home.config(state="disabled")
            QuickSNAP.config(state="disabled")
            docx_list=[]
            # import glob
            # def clean_document():
            #     print("cleaning document..")
            #     tail_dir=os.getcwd()
            #     nonlocal docx_list
            #     try:
            #         os.chdir(r"{}\document".format(path))
            #         docx_file="*"
            #         docx_list= glob.glob(docx_file)
            #         for i in docx_list:
            #             os.remove(i)
            #         os.chdir(tail_dir)
            #     except PermissionError:
            #         home.config(state="normal")
            #         QuickSNAP.config(state="normal")
            #         start.config(state='normal', bg=bg_button["color"])
            #         stop.config(state='disabled', bg=bg_button["color"])
            #         tmsg.showwarning("Warning",
            #                          "Not able to create, Permission denied.\nPlease close document if open\n{}".format(
            #                              sys.exc_info()), parent=window1)
            #
            #     except:
            #         os.mkdir(path+"\document")

            value = tmsg.askquestion("Warning", "Do you want to split into mutilple document?\nSplit the document into multiple document if size exceeds 1 MB", parent=window1)
            # clean_document()
            try:
                image_list, actual_path=create_document(document_name, path)
                name = document_name
                global metadata
                length = len(image_list)
                count = 0
                document = Document()
                section = document.sections[0]
                section.page_height = Mm(297)
                section.page_width = Mm(210)
                section.left_margin = Mm(25.4)
                section.right_margin = Mm(25.4)
                section.top_margin = Mm(25.4)
                section.bottom_margin = Mm(25.4)
                section.header_distance = Mm(12.7)
                section.footer_distance = Mm(12.7)

                document.add_paragraph(metadata)
                if length > 0:
                    def append_image():
                        nonlocal count, name, document
                        try:
                            document.save("{}\document\{}.docx".format(path, name))
                            if value == "no":
                                document.add_picture("{}\image\{}".format(path, image_list[count]), width=Inches(6.27),
                                                     height=Inches(3.52))
                                document.save("{}\document\{}.docx".format(path, document_name))
                                message_show("Processing {}/{} ".format(count + 1, length))
                                count += 1
                            else:
                                if os.path.getsize("{}\document\{}.docx".format(path, name)) < 900000:
                                    document.add_picture("{}\image\{}".format(path, image_list[count]),
                                                         width=Inches(6.27), height=Inches(3.52))
                                    document.save("{}\document\{}.docx".format(path, name))
                                    message_show("Processing {}/{} ".format(count + 1, length))
                                    count += 1
                                else:
                                    name = "split_" + str(count)
                                    print(name)
                                    document = Document()
                                    document.add_picture("{}\image\{}".format(path, image_list[count]),
                                                         width=Inches(6.27), height=Inches(3.52))
                                    document.save("{}\document\{}.docx".format(path, name))
                                    message_show("Processing {}/{} ".format(count + 1, length))
                                    count = count + 1

                            if count == length:
                                message_show("Document created on {}".format(datetime.now().strftime('%H:%M:%S')))
                                save_button.config(state='normal', bg=bg_button["color"])
                                start.config(state='normal', bg=bg_button["color"])
                                stop.config(state='disabled', bg=bg_button["color"])
                                home.config(state="normal")
                                QuickSNAP.config(state="normal")
                                os.chdir(r"{}".format(actual_path))
                                tmsg.showinfo("Done",
                                              "Document created on {}".format(datetime.now().strftime('%H:%M:%S')),
                                              parent=window1)
                                os.startfile(path + "\\document")
                            else:
                                window1.after_idle(append_image)

                        except PermissionError:
                            os.chdir(r"{}".format(actual_path))
                            message_show("{}-\nNot able to create, permission denied.\nclose document".format(
                                datetime.now().strftime('%H:%M:%S')))
                            save_button.config(state='normal', bg=bg_button["color"])
                            home.config(state="normal")
                            QuickSNAP.config(state="normal")
                            start.config(state='normal', bg=bg_button["color"])
                            stop.config(state='disabled', bg=bg_button["color"])
                            tmsg.showerror("Warning",
                                           "Not able to create, Permission denied.\nPlease close document if open\n{}".format(
                                               sys.exc_info()), parent=window1)

                        except FileNotFoundError:
                            os.chdir(r"{}".format(actual_path))
                            message_show(
                                "{}-\nNot able to create, Total length of document name\n and output directory path\nlength should not exceed 260 characters".format(
                                    datetime.now().strftime('%H:%M:%S')))
                            save_button.config(state='normal', bg=bg_button["color"])
                            home.config(state="normal")
                            QuickSNAP.config(state="normal")
                            start.config(state='normal', bg=bg_button["color"])
                            stop.config(state='disabled', bg=bg_button["color"])
                            tmsg.showerror("Warning",
                                           "Not able to create, Total length of document name\nand output directory path length should\nnot exceed 260 characters\n{}".format(
                                               sys.exc_info()), parent=window1)

                        except Exception as e:
                            os.chdir(r"{}".format(actual_path))
                            message_show("{}-\nNot able to create".format(datetime.now().strftime('%H:%M:%S')))
                            save_button.config(state='normal', bg=bg_button["color"])
                            home.config(state="normal")
                            QuickSNAP.config(state="normal")
                            start.config(state='normal', bg=bg_button["color"])
                            stop.config(state='disabled', bg=bg_button["color"])
                            tmsg.showerror("Warning", e, parent=window1)
                    window1.after_idle(append_image)
                else:
                    try:
                        document.add_paragraph("No image to attach")
                        document.save("{}\document\{}.docx".format(path, document_name))
                        message_show("Document created with no image attachment")
                        save_button.config(state='normal', bg=bg_button["color"])
                        home.config(state="normal")
                        QuickSNAP.config(state="normal")
                        start.config(state='normal', bg=bg_button["color"])
                        stop.config(state='disabled', bg=bg_button["color"])
                        os.chdir(r"{}".format(actual_path))
                    except FileNotFoundError:
                        os.chdir(r"{}".format(actual_path))
                        message_show(
                            "{}-\nNot able to create, Total length of document name\n and output directory path length should not\nexceed 260 characters".format(
                                datetime.now().strftime('%H:%M:%S')))
                        save_button.config(state='normal', bg=bg_button["color"])
                        home.config(state="normal")
                        QuickSNAP.config(state="normal")
                        start.config(state='normal', bg=bg_button["color"])
                        stop.config(state='disabled', bg=bg_button["color"])
                        tmsg.showerror("Warning",
                                       "Not able to create, Total length of document\nname and output directory path length should not\nexceed 260 characters\n{}".format(
                                           sys.exc_info()), parent=window1)
                    except Exception as e:
                        os.chdir(r"{}".format(actual_path))
                        message_show("{}-\nNot able to create".format(datetime.now().strftime('%H:%M:%S')))
                        save_button.config(state='normal', bg=bg_button["color"])
                        home.config(state="normal")
                        QuickSNAP.config(state="normal")
                        start.config(state='normal', bg=bg_button["color"])
                        stop.config(state='disabled', bg=bg_button["color"])
                        tmsg.showerror("Warning", e, parent=window1)
            except Exception as e:
                tmsg.showerror("Warning",e, parent=window1)
                save_button.config(state='normal', bg=bg_button["color"])
                home.config(state="normal")
                QuickSNAP.config(state="normal")
                start.config(state='normal', bg=bg_button["color"])
                stop.config(state='disabled', bg=bg_button["color"])
            # name=document_name
            # global metadata
            # length=len(image_list)
            # count=0
            # document = Document()
            # section = document.sections[0]
            # section.page_height = Mm(297)
            # section.page_width = Mm(210)
            # section.left_margin = Mm(25.4)
            # section.right_margin = Mm(25.4)
            # section.top_margin = Mm(25.4)
            # section.bottom_margin = Mm(25.4)
            # section.header_distance = Mm(12.7)
            # section.footer_distance = Mm(12.7)
            #
            # document.add_paragraph(metadata)
            # if length >0:
            #     def append_image():
            #         nonlocal count, name, document
            #         try:
            #             document.save("{}\document\{}.docx".format(path, name))
            #             if value == "no":
            #                 document.add_picture("{}\image\{}".format(path,image_list[count]), width=Inches(6.27), height=Inches(3.52))
            #                 document.save("{}\document\{}.docx".format(path,document_name))
            #                 message_show("Processing {}/{} ".format(count+1, length))
            #                 count+=1
            #             else:
            #                 if os.path.getsize("{}\document\{}.docx".format(path,name)) < 900000:
            #                     document.add_picture("{}\image\{}".format(path,image_list[count]), width=Inches(6.27), height=Inches(3.52))
            #                     document.save("{}\document\{}.docx".format(path,name))
            #                     message_show("Processing {}/{} ".format(count + 1, length))
            #                     count+=1
            #                 else:
            #                     name="split_"+str(count)
            #                     print(name)
            #                     document = Document()
            #                     document.add_picture("{}\image\{}".format(path,image_list[count]), width=Inches(6.27), height=Inches(3.52))
            #                     document.save("{}\document\{}.docx".format(path,name))
            #                     message_show("Processing {}/{} ".format(count + 1, length))
            #                     count = count + 1
            #
            #             if count == length:
            #                 message_show("Document created on {}".format(datetime.now().strftime('%H:%M:%S')))
            #                 save_button.config(state='normal', bg=bg_button["color"])
            #                 start.config(state='normal', bg=bg_button["color"])
            #                 stop.config(state='disabled', bg=bg_button["color"])
            #                 home.config(state="normal")
            #                 QuickSNAP.config(state="normal")
            #                 os.chdir(r"{}".format(actual_path))
            #                 tmsg.showinfo("Done","Document created on {}".format(datetime.now().strftime('%H:%M:%S')), parent=window1)
            #                 os.startfile(path + "\\document")
            #             else:
            #                 window1.after_idle(append_image)
            #
            #         except PermissionError:
            #             os.chdir(r"{}".format(actual_path))
            #             message_show("{}-\nNot able to create, permission denied.\nclose document".format(datetime.now().strftime('%H:%M:%S')))
            #             save_button.config(state='normal', bg=bg_button["color"])
            #             home.config(state="normal")
            #             QuickSNAP.config(state="normal")
            #             start.config(state='normal', bg=bg_button["color"])
            #             stop.config(state='disabled', bg=bg_button["color"])
            #             tmsg.showerror("Warning", "Not able to create, Permission denied.\nPlease close document if open\n{}".format(sys.exc_info()), parent=window1)
            #
            #         except FileNotFoundError:
            #             os.chdir(r"{}".format(actual_path))
            #             message_show("{}-\nNot able to create, Total length of document name\n and output directory path\nlength should not exceed 260 characters".format(datetime.now().strftime('%H:%M:%S')))
            #             save_button.config(state='normal', bg=bg_button["color"])
            #             home.config(state="normal")
            #             QuickSNAP.config(state="normal")
            #             start.config(state='normal', bg=bg_button["color"])
            #             stop.config(state='disabled', bg=bg_button["color"])
            #             tmsg.showerror("Warning","Not able to create, Total length of document name\nand output directory path length should\nsnot exceed 260 characters\n{}".format(sys.exc_info()), parent=window1)
            #
            #         except Exception as e:
            #             os.chdir(r"{}".format(actual_path))
            #             message_show("{}-\nNot able to create".format(datetime.now().strftime('%H:%M:%S')))
            #             save_button.config(state='normal', bg=bg_button["color"])
            #             home.config(state="normal")
            #             QuickSNAP.config(state="normal")
            #             start.config(state='normal', bg=bg_button["color"])
            #             stop.config(state='disabled', bg=bg_button["color"])
            #             tmsg.showerror("Warning",e,parent=window1)
            #
            #     window1.after_idle(append_image)
            #
            # else:
            #     try:
            #         document.add_paragraph("No image to attach")
            #         document.save("{}\document\{}.docx".format(path, document_name))
            #         message_show("Document created with no image attachment")
            #         save_button.config(state='normal',bg=bg_button["color"])
            #         home.config(state="normal")
            #         QuickSNAP.config(state="normal")
            #         start.config(state='normal', bg=bg_button["color"])
            #         stop.config(state='disabled', bg=bg_button["color"])
            #         os.chdir(r"{}".format(actual_path))
            #     except FileNotFoundError:
            #         os.chdir(r"{}".format(actual_path))
            #         message_show(
            #             "{}-\nNot able to create, Total length of document name\n and output directory path length should not\nexceed 260 characters".format(
            #                 datetime.now().strftime('%H:%M:%S')))
            #         save_button.config(state='normal', bg=bg_button["color"])
            #         home.config(state="normal")
            #         QuickSNAP.config(state="normal")
            #         start.config(state='normal', bg=bg_button["color"])
            #         stop.config(state='disabled', bg=bg_button["color"])
            #         tmsg.showerror("Warning","Not able to create, Total length of document\nname and output directory path length should not\nexceed 260 characters\n{}".format(sys.exc_info()), parent=window1)
            #
            #     except Exception as e:
            #         os.chdir(r"{}".format(actual_path))
            #         message_show("{}-\nNot able to create".format(datetime.now().strftime('%H:%M:%S')))
            #         save_button.config(state='normal', bg=bg_button["color"])
            #         home.config(state="normal")
            #         QuickSNAP.config(state="normal")
            #         start.config(state='normal', bg=bg_button["color"])
            #         stop.config(state='disabled', bg=bg_button["color"])
            #         tmsg.showerror("Warning", e, parent=window1)
        # ******************* Notification print
        f7=Frame(f1, relief=SUNKEN, bg=color["frame_bg"],highlightbackground=color["highlightbackground"], highlightcolor=color["highlightcolor"])
        f7.pack(fill='both')
        global message_label, start_icon, save_icon, save_icon
        message_label = Label(f7, text="No new notifications",font=('Microsoft Sans Serif', 10) , justify=LEFT, fg="#990033", bg="white")
        message_label.grid(sticky=W, pady=5)


        f8 = Frame(f1, relief=SUNKEN)
        start = Button(f8, text='START', image=start_icon,font="Balthazar 12 bold",compound="left", bd=0.4,bg=bg_button["color"],fg="black", command=thread_create)
        stop = Button(f8, text='STOP',image=stop_icon, font="Balthazar 12 bold",compound="left",   bd=0.4,fg="black", state='disabled', bg=bg_button["color"],
                      command=quit_)
        global save_button
        save_button = Button(f8, text='SAVE',image=save_icon, font="Balthazar 12 bold",compound="left" , bd=0.4, bg=bg_button["color"],fg="black", command=save)

        f8.pack(anchor='center', side=BOTTOM, pady=5)
        start.pack(anchor='center', side="left")
        stop.pack(anchor='center', side="left", padx=2)
        save_button.pack(anchor='center', side="left", padx=2)


    def statupScrenshot():
        myScreenshot = pyautogui.screenshot()
        myScreenshot.save(os.getcwd()+"\\startup.png")


    # **************** Bottom Section *********************
    def call():
        value = tmsg.askquestion("Warning", "Are you done with the screenshot?", parent=window1)
        if value == "yes":
            window1.withdraw()
            root.deiconify()

    def bottom_section():
        def open_file_exp(flag):
            global path
            try:
                if flag == 1:
                    os.startfile(path+"\\image")
                elif flag == 0:
                    os.startfile(path+"\\document")
            except:
                if flag == 1:
                    try:
                        os.mkdir(path+"\\image")
                        os.startfile(path+"\\image")
                    except Exception as e:
                        tmsg.showerror("Warning", e, parent=window1)
                elif flag == 0:
                    try:
                        os.mkdir(path + "\\document")
                        os.startfile(path+"\\document")
                    except Exception as e:
                        tmsg.showerror("Warning", e, parent=window1)



        def open_window(path, window1):
            window1.withdraw()
            floating_window(window1, path)


        ''' Button for home'''
        global home
        font = {"font": "Balthazar"}
        home=Button(f4, image=close_photo,text="Home", font=(font['font'], 7, "bold") ,command=call,compound="left")
        home.grid(row=0)

        btn3 = Button(f4, image=imagefolder_icon,text="ImageFolder",font=(font['font'], 7, "bold") ,
                      command=lambda : open_file_exp(1), compound="left")


        btn3.grid(row=0, column=2)

        btn4 = Button(f4,font=(font['font'], 7, "bold"), image=docsfolder_icon,text="DocsFolder",
                      command=lambda : open_file_exp(0), compound="left")
        btn4.grid(row=0, column=3)
        global QuickSNAP
        QuickSNAP = Button(f4,font=(font['font'], 7, "bold"), image=screenshot_icon,text="QuickSNAP",
                      command=lambda: open_window(path, window1), compound="left")
        QuickSNAP.grid(row=0, column=4)




        # *************** Middle Section ***********************
    def right_section():
        image_label = Label(f3, image=photo, font="Helvetica 16 bold", fg="red", pady=22)
        image_label.pack()
        bottom_section()



    statupScrenshot()
    read_metadata()
    image_handle()
    headerSection()
    side_section()
    right_section()
    bottom_section()
    window1.protocol("WM_DELETE_WINDOW", root.destroy)
    root.protocol("WM_DELETE_WINDOW", root.destroy)
    window1.mainloop()
